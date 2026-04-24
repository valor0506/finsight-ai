"""
report_builder/docx_builder.py
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from datetime import datetime
import io

DARK_BG    = "0D1117"
ACCENT     = "00C896"
SUBACCENT  = "1A73E8"
LIGHT_GREY = "F4F6F9"
MID_GREY   = "6B7280"
WHITE      = "FFFFFF"


def _set_cell_bg(cell, hex_color: str):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)


def _add_run(para, text: str, bold=False, italic=False,
             size_pt=11, color=None, font="Calibri"):
    run = para.add_run(text)
    run.bold        = bold
    run.italic      = italic
    run.font.name   = font
    run.font.size   = Pt(size_pt)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    return run


def _add_divider(doc, color=ACCENT, thickness=8):
    p   = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pb  = OxmlElement("w:pBdr")
    bot = OxmlElement("w:bottom")
    bot.set(qn("w:val"),   "single")
    bot.set(qn("w:sz"),    str(thickness))
    bot.set(qn("w:space"), "1")
    bot.set(qn("w:color"), color)
    pb.append(bot)
    pPr.append(pb)
    p.paragraph_format.space_after  = Pt(4)
    p.paragraph_format.space_before = Pt(4)
    return p


def _kv_table(doc, rows: list, col_widths=(5.5, 4.0)):
    table = doc.add_table(rows=len(rows), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    for i, (key, val) in enumerate(rows):
        row   = table.rows[i]
        bg    = LIGHT_GREY if i % 2 == 0 else WHITE
        left  = row.cells[0]
        right = row.cells[1]
        _set_cell_bg(left,  bg)
        _set_cell_bg(right, bg)
        left.width  = Inches(col_widths[0])
        right.width = Inches(col_widths[1])
        _add_run(left.paragraphs[0],  str(key),  bold=True, size_pt=10, color="374151")
        _add_run(right.paragraphs[0], str(val) if val is not None else "N/A", size_pt=10)
    doc.add_paragraph()


def _section_heading(doc, title: str):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(2)
    _add_run(p, f"  {title}", bold=True, size_pt=13, color=WHITE, font="Calibri")
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  "1B2A3B")
    pPr.append(shd)
    return p


def _parse_analysis_text(doc, text: str):
    for line in text.splitlines():
        line = line.strip()
        if not line:
            doc.add_paragraph()
            continue
        if line.startswith("## "):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(10)
            _add_run(p, line[3:], bold=True, size_pt=12, color=SUBACCENT)
        elif line.startswith("# "):
            _section_heading(doc, line[2:])
        elif line.startswith("- ") or line.startswith("* "):
            p = doc.add_paragraph(style="List Bullet")
            _add_run(p, line[2:], size_pt=10.5)
        else:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(4)
            parts = line.split("**")
            for idx, part in enumerate(parts):
                if part:
                    _add_run(p, part, bold=(idx % 2 == 1), size_pt=10.5)


def _build_cover(doc, symbol: str, report_type: str = "Commodity"):
    header_p = doc.add_paragraph()
    pPr = header_p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto"); shd.set(qn("w:fill"), DARK_BG)
    pPr.append(shd)
    _add_run(header_p, f"FinSight AI  |  {symbol} {report_type} Intelligence Report",
             bold=True, size_pt=20, color=ACCENT, font="Calibri")

    sub_p = doc.add_paragraph()
    pPr2 = sub_p._p.get_or_add_pPr()
    shd2 = OxmlElement("w:shd")
    shd2.set(qn("w:val"), "clear"); shd2.set(qn("w:color"), "auto"); shd2.set(qn("w:fill"), DARK_BG)
    pPr2.append(shd2)
    ts = datetime.utcnow().strftime("%d %b %Y, %H:%M UTC")
    _add_run(sub_p,
             f"Generated: {ts}  |  Powered by Gemini 1.5 Flash  |  Data: Alpha Vantage + FRED + NewsAPI",
             size_pt=9, color=MID_GREY)
    doc.add_paragraph()


def _build_commodity_snapshot(doc, symbol: str, data: dict, macro: dict):
    _section_heading(doc, "Market Snapshot")
    doc.add_paragraph()
    rows = [
        ("Current Price",       f"{data.get('current_price')} {data.get('currency', 'USD')}"),
        ("52-Week High",        f"{data.get('week_52_high')} USD"),
        ("52-Week Low",         f"{data.get('week_52_low')} USD"),
        ("% from 52W High",     f"{data.get('pct_from_52w_high')}%"),
        ("RSI-14",              str(data.get("rsi_14"))),
        ("SMA-20",              str(data.get("sma_20"))),
        ("SMA-50",              str(data.get("sma_50"))),
        ("── Macro Context ──", ""),
        ("DXY (USD Index)",     str(macro.get("DXY", {}).get("value"))),
        ("US 10Y Yield",        f"{macro.get('US_10Y_YIELD', {}).get('value')}%"),
        ("Gold-Silver Ratio",   str(macro.get("GOLD_SILVER_RATIO", {}).get("value"))),
        ("USD/INR",             str(macro.get("USDINR", {}).get("value"))),
    ]
    _kv_table(doc, rows)


def _build_equity_snapshot(doc, ticker: str, data: dict, macro: dict):
    _section_heading(doc, "Market Snapshot")
    doc.add_paragraph()
    rows = [
        ("Ticker",            ticker),
        ("Current Price",     f"{data.get('current_price')} {data.get('currency', 'USD')}"),
        ("52-Week High",      str(data.get("week_52_high"))),
        ("52-Week Low",       str(data.get("week_52_low"))),
        ("RSI-14",            str(data.get("rsi_14"))),
        ("SMA-50",            str(data.get("sma_50"))),
        ("SMA-200",           str(data.get("sma_200"))),
        ("── Macro Context ──", ""),
        ("DXY (USD Index)",   str(macro.get("DXY", {}).get("value"))),
        ("US 10Y Yield",      f"{macro.get('US_10Y_YIELD', {}).get('value')}%"),
        ("VIX",               str(macro.get("VIX_US", {}).get("value"))),
        ("USD/INR",           str(macro.get("USDINR", {}).get("value"))),
    ]
    _kv_table(doc, rows)


def _base_doc() -> Document:
    doc = Document()
    for section in doc.sections:
        section.top_margin    = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin   = Inches(1.1)
        section.right_margin  = Inches(1.1)
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    return doc


def _add_disclaimer(doc):
    _add_divider(doc, color=MID_GREY, thickness=4)
    footer_p = doc.add_paragraph()
    footer_p.paragraph_format.space_before = Pt(8)
    _add_run(
        footer_p,
        "DISCLAIMER: This report is AI-generated for informational purposes only. "
        "Not financial advice. Always consult a qualified financial advisor before making investment decisions.",
        size_pt=8, italic=True, color=MID_GREY,
    )


def _to_bytes(doc: Document) -> bytes:
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


# ── Public functions ───────────────────────────────────────────

def build_commodity_report(symbol: str, data: dict, analysis_text: str, macro: dict) -> bytes:
    doc = _base_doc()
    _build_cover(doc, symbol, "Commodity")
    _add_divider(doc, color=ACCENT)
    _build_commodity_snapshot(doc, symbol, data, macro)
    _add_divider(doc, color=ACCENT)
    _section_heading(doc, "AI-Generated Analysis")
    doc.add_paragraph()
    _parse_analysis_text(doc, analysis_text)
    _add_disclaimer(doc)
    return _to_bytes(doc)


def build_equity_report(ticker: str, data: dict, analysis_text: str, macro: dict) -> bytes:
    doc = _base_doc()
    _build_cover(doc, ticker, "Equity")
    _add_divider(doc, color=ACCENT)
    _build_equity_snapshot(doc, ticker, data, macro)
    _add_divider(doc, color=ACCENT)
    _section_heading(doc, "AI-Generated Analysis")
    doc.add_paragraph()
    _parse_analysis_text(doc, analysis_text)
    _add_disclaimer(doc)
    return _to_bytes(doc)